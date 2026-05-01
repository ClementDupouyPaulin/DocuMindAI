from pathlib import Path

from docx import Document as DocxDocument

from app.services.extraction_service import extract_text_from_docx, extract_text_from_txt


def test_extract_text_from_txt(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Hello DocuMind AI", encoding="utf-8")

    result = extract_text_from_txt(file_path)

    assert len(result) == 1
    assert result[0]["page_number"] is None
    assert result[0]["text"] == "Hello DocuMind AI"


def test_extract_text_from_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"

    document = DocxDocument()
    document.add_paragraph("Hello DocuMind AI")
    document.add_paragraph("This is a DOCX extraction test.")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Column A"
    table.cell(0, 1).text = "Column B"

    document.save(file_path)

    result = extract_text_from_docx(file_path)

    assert len(result) == 1
    assert result[0]["page_number"] is None
    assert "Hello DocuMind AI" in str(result[0]["text"])
    assert "This is a DOCX extraction test." in str(result[0]["text"])
    assert "Column A | Column B" in str(result[0]["text"])
